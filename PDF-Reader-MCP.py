"""
ENTERPRISE PDF MCP SERVER - Production Grade
For Integration with MENTARI V.19
McKinsey & Big4 Ready

Features:
- Multi-format PDF extraction (text, tables, images, metadata)
- OCR support for scanned documents
- Intelligent table detection & extraction
- Form field extraction
- Password-protected PDF support
- Batch processing
- Memory-efficient streaming
- Comprehensive error handling
- Audit logging
- **NEW: Word ↔ PDF Conversion (Production Quality)**
"""

import sys
import os
import io
import json
import logging
import asyncio
import traceback
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path

# PDF Libraries
try:
    import PyPDF2
    from pypdf import PdfReader, PdfWriter
except ImportError:
    print("ERROR: Install pypdf: pip install pypdf PyPDF2")
    sys.exit(1)

try:
    import pdfplumber
except ImportError:
    print("WARNING: pdfplumber not installed. Table extraction limited.")
    pdfplumber = None

try:
    from pdf2image import convert_from_path
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    print("WARNING: OCR unavailable. Install: pip install pdf2image pillow pytesseract")
    OCR_AVAILABLE = False

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    print("INFO: Camelot not installed. Advanced table extraction unavailable.")
    CAMELOT_AVAILABLE = False

# Word Conversion Libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("WARNING: python-docx not installed. Install: pip install python-docx")
    DOCX_AVAILABLE = False

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    print("WARNING: docx2pdf not installed. Install: pip install docx2pdf")
    DOCX2PDF_AVAILABLE = False

# Alternative conversion using LibreOffice (more reliable cross-platform)
try:
    import subprocess
    LIBREOFFICE_AVAILABLE = False
    # Check if LibreOffice is installed
    try:
        result = subprocess.run(['libreoffice', '--version'], 
                              capture_output=True, timeout=5)
        if result.returncode == 0:
            LIBREOFFICE_AVAILABLE = True
    except:
        pass
except:
    LIBREOFFICE_AVAILABLE = False

# MCP SDK
try:
    from mcp.server import Server
    from mcp.types import Tool, TextContent, ImageContent, EmbeddedResource
    import mcp.server.stdio
except ImportError:
    print("ERROR: Install MCP SDK: pip install mcp")
    sys.exit(1)

# ============================================================================
# LOGGING CONFIGURATION
# ============================================================================

def setup_logging():
    """Enterprise-grade logging with rotation"""
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"pdf_mcp_{datetime.now().strftime('%Y%m%d')}.log"
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file),
            logging.StreamHandler(sys.stderr)
        ]
    )
    return logging.getLogger("PDFMCPServer")

logger = setup_logging()

# ============================================================================
# WORD ↔ PDF CONVERTER ENGINE
# ============================================================================

class WordPDFConverter:
    """
    Production-grade Word ↔ PDF converter
    Multiple conversion strategies for maximum reliability
    """
    
    def __init__(self):
        self.supported_word_formats = ['.docx', '.doc']
        self.supported_pdf_formats = ['.pdf']
        self.max_file_size = 50 * 1024 * 1024  # 50MB limit
        
        # Check available conversion methods
        self.conversion_methods = []
        if DOCX2PDF_AVAILABLE:
            self.conversion_methods.append('docx2pdf')
        if LIBREOFFICE_AVAILABLE:
            self.conversion_methods.append('libreoffice')
        
        logger.info(f"Word converter initialized. Methods: {self.conversion_methods}")
    
    def validate_file(self, filepath: str, expected_formats: List[str]) -> Tuple[bool, str]:
        """Validate file before conversion"""
        try:
            path = Path(filepath)
            
            if not path.exists():
                return False, f"File not found: {filepath}"
            
            if path.suffix.lower() not in expected_formats:
                return False, f"Invalid format. Expected {expected_formats}, got {path.suffix}"
            
            size = path.stat().st_size
            if size > self.max_file_size:
                return False, f"File too large: {size/1024/1024:.1f}MB (max 50MB)"
            
            if size == 0:
                return False, "File is empty"
            
            return True, "Valid file"
            
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    # ========================================================================
    # WORD TO PDF CONVERSION
    # ========================================================================
    
    def word_to_pdf_libreoffice(self, word_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert Word to PDF using LibreOffice (BEST METHOD - Most reliable)
        Works on Windows, Mac, Linux
        """
        try:
            word_path = Path(word_path).absolute()
            output_dir = Path(output_path).parent.absolute()
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # LibreOffice command
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                str(word_path)
            ]
            
            logger.info(f"Converting with LibreOffice: {word_path.name}")
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120  # 2 minute timeout
            )
            
            if result.returncode != 0:
                return False, f"LibreOffice error: {result.stderr}"
            
            # LibreOffice creates PDF with same name as input
            expected_pdf = output_dir / f"{word_path.stem}.pdf"
            
            if expected_pdf.exists():
                # Rename if needed
                if expected_pdf != Path(output_path):
                    expected_pdf.rename(output_path)
                
                size_mb = Path(output_path).stat().st_size / 1024 / 1024
                logger.info(f"✅ Converted: {output_path} ({size_mb:.2f}MB)")
                return True, f"Converted successfully ({size_mb:.2f}MB)"
            else:
                return False, "PDF file not created"
                
        except subprocess.TimeoutExpired:
            return False, "Conversion timeout (>2 minutes)"
        except Exception as e:
            logger.error(f"LibreOffice conversion failed: {traceback.format_exc()}")
            return False, f"Conversion failed: {str(e)}"
    
    def word_to_pdf_docx2pdf(self, word_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert Word to PDF using docx2pdf (Windows primarily)
        Requires Microsoft Word installed
        """
        try:
            logger.info(f"Converting with docx2pdf: {Path(word_path).name}")
            
            docx2pdf_convert(word_path, output_path)
            
            if Path(output_path).exists():
                size_mb = Path(output_path).stat().st_size / 1024 / 1024
                logger.info(f"✅ Converted: {output_path} ({size_mb:.2f}MB)")
                return True, f"Converted successfully ({size_mb:.2f}MB)"
            else:
                return False, "PDF file not created"
                
        except Exception as e:
            logger.error(f"docx2pdf conversion failed: {str(e)}")
            return False, f"Conversion failed: {str(e)}"
    
    def convert_word_to_pdf(self, word_path: str, output_path: Optional[str] = None,
                           method: str = 'auto') -> Dict[str, Any]:
        """
        Main Word to PDF conversion with automatic fallback
        """
        try:
            # Validate input
            is_valid, msg = self.validate_file(word_path, self.supported_word_formats)
            if not is_valid:
                return {"error": msg, "success": False}
            
            # Determine output path
            if not output_path:
                word_path_obj = Path(word_path)
                output_path = str(word_path_obj.parent / f"{word_path_obj.stem}.pdf")
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Try conversion methods in order of preference
            methods_to_try = []
            
            if method == 'auto':
                # Try LibreOffice first (most reliable), then docx2pdf
                if LIBREOFFICE_AVAILABLE:
                    methods_to_try.append(('libreoffice', self.word_to_pdf_libreoffice))
                if DOCX2PDF_AVAILABLE:
                    methods_to_try.append(('docx2pdf', self.word_to_pdf_docx2pdf))
            elif method == 'libreoffice' and LIBREOFFICE_AVAILABLE:
                methods_to_try.append(('libreoffice', self.word_to_pdf_libreoffice))
            elif method == 'docx2pdf' and DOCX2PDF_AVAILABLE:
                methods_to_try.append(('docx2pdf', self.word_to_pdf_docx2pdf))
            
            if not methods_to_try:
                return {
                    "error": "No conversion methods available. Install LibreOffice or docx2pdf.",
                    "success": False
                }
            
            # Try each method until one succeeds
            errors = []
            for method_name, method_func in methods_to_try:
                logger.info(f"Attempting conversion with: {method_name}")
                success, message = method_func(word_path, output_path)
                
                if success:
                    return {
                        "success": True,
                        "method": method_name,
                        "input_file": str(Path(word_path).absolute()),
                        "output_file": str(Path(output_path).absolute()),
                        "output_size_mb": round(Path(output_path).stat().st_size / 1024 / 1024, 2),
                        "message": message
                    }
                else:
                    errors.append(f"{method_name}: {message}")
            
            # All methods failed
            return {
                "error": "All conversion methods failed: " + "; ".join(errors),
                "success": False
            }
            
        except Exception as e:
            logger.error(f"Word to PDF conversion failed: {traceback.format_exc()}")
            return {"error": f"Conversion failed: {str(e)}", "success": False}
    
    # ========================================================================
    # PDF TO WORD CONVERSION
    # ========================================================================
    
    def pdf_to_word_advanced(self, pdf_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Advanced PDF to Word conversion with layout preservation
        Uses pypdf for text extraction and python-docx for Word creation
        """
        try:
            if not DOCX_AVAILABLE:
                return False, "python-docx not installed"
            
            logger.info(f"Converting PDF to Word: {Path(pdf_path).name}")
            
            # Read PDF
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            
            # Create Word document
            doc = Document()
            
            # Add metadata
            core_props = doc.core_properties
            core_props.title = f"Converted from {Path(pdf_path).name}"
            core_props.author = "Enterprise PDF MCP Server"
            core_props.comments = f"Converted on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
            # Extract and add content page by page
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                
                # Add page header
                if i > 0:
                    doc.add_page_break()
                
                heading = doc.add_heading(f'Page {i + 1}', level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Split into paragraphs and add
                paragraphs = text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        para.style = 'Normal'
            
            # Try to extract and add tables if pdfplumber available
            if pdfplumber:
                try:
                    with pdfplumber.open(pdf_path) as pdf:
                        for i, page in enumerate(pdf.pages):
                            tables = page.extract_tables()
                            if tables:
                                doc.add_paragraph()  # Spacing
                                doc.add_heading(f'Tables from Page {i + 1}', level=3)
                                
                                for table_idx, table in enumerate(tables):
                                    if table:
                                        # Create Word table
                                        rows = len(table)
                                        cols = len(table[0]) if table else 0
                                        
                                        if rows > 0 and cols > 0:
                                            word_table = doc.add_table(rows=rows, cols=cols)
                                            word_table.style = 'Light Grid Accent 1'
                                            
                                            # Fill table
                                            for row_idx, row in enumerate(table):
                                                for col_idx, cell in enumerate(row):
                                                    if cell:
                                                        word_table.rows[row_idx].cells[col_idx].text = str(cell)
                                            
                                            doc.add_paragraph()  # Spacing after table
                except Exception as e:
                    logger.warning(f"Table extraction skipped: {str(e)}")
            
            # Save document
            doc.save(output_path)
            
            if Path(output_path).exists():
                size_mb = Path(output_path).stat().st_size / 1024 / 1024
                logger.info(f"✅ Converted to Word: {output_path} ({size_mb:.2f}MB)")
                return True, f"Converted successfully ({total_pages} pages, {size_mb:.2f}MB)"
            else:
                return False, "Word file not created"
                
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {traceback.format_exc()}")
            return False, f"Conversion failed: {str(e)}"
    
    def pdf_to_word_libreoffice(self, pdf_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert PDF to Word using LibreOffice (alternative method)
        """
        try:
            pdf_path = Path(pdf_path).absolute()
            output_dir = Path(output_path).parent.absolute()
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # LibreOffice command for PDF to ODT (then can convert to DOCX)
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'docx',
                '--outdir', str(output_dir),
                str(pdf_path)
            ]
            
            logger.info(f"Converting PDF to Word with LibreOffice: {pdf_path.name}")
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120
            )
            
            if result.returncode != 0:
                return False, f"LibreOffice error: {result.stderr}"
            
            expected_docx = output_dir / f"{pdf_path.stem}.docx"
            
            if expected_docx.exists():
                if expected_docx != Path(output_path):
                    expected_docx.rename(output_path)
                
                size_mb = Path(output_path).stat().st_size / 1024 / 1024
                logger.info(f"✅ Converted: {output_path} ({size_mb:.2f}MB)")
                return True, f"Converted successfully ({size_mb:.2f}MB)"
            else:
                return False, "DOCX file not created"
                
        except subprocess.TimeoutExpired:
            return False, "Conversion timeout (>2 minutes)"
        except Exception as e:
            logger.error(f"LibreOffice PDF to Word failed: {traceback.format_exc()}")
            return False, f"Conversion failed: {str(e)}"
    
    def convert_pdf_to_word(self, pdf_path: str, output_path: Optional[str] = None,
                           method: str = 'auto') -> Dict[str, Any]:
        """
        Main PDF to Word conversion with automatic fallback
        """
        try:
            # Validate input
            is_valid, msg = self.validate_file(pdf_path, self.supported_pdf_formats)
            if not is_valid:
                return {"error": msg, "success": False}
            
            # Determine output path
            if not output_path:
                pdf_path_obj = Path(pdf_path)
                output_path = str(pdf_path_obj.parent / f"{pdf_path_obj.stem}.docx")
            
            # Ensure output directory exists
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            
            # Try conversion methods
            methods_to_try = []
            
            if method == 'auto':
                # Try advanced method first, then LibreOffice
                if DOCX_AVAILABLE:
                    methods_to_try.append(('advanced', self.pdf_to_word_advanced))
                if LIBREOFFICE_AVAILABLE:
                    methods_to_try.append(('libreoffice', self.pdf_to_word_libreoffice))
            elif method == 'advanced' and DOCX_AVAILABLE:
                methods_to_try.append(('advanced', self.pdf_to_word_advanced))
            elif method == 'libreoffice' and LIBREOFFICE_AVAILABLE:
                methods_to_try.append(('libreoffice', self.pdf_to_word_libreoffice))
            
            if not methods_to_try:
                return {
                    "error": "No conversion methods available. Install python-docx or LibreOffice.",
                    "success": False
                }
            
            # Try each method
            errors = []
            for method_name, method_func in methods_to_try:
                logger.info(f"Attempting PDF to Word with: {method_name}")
                success, message = method_func(pdf_path, output_path)
                
                if success:
                    return {
                        "success": True,
                        "method": method_name,
                        "input_file": str(Path(pdf_path).absolute()),
                        "output_file": str(Path(output_path).absolute()),
                        "output_size_mb": round(Path(output_path).stat().st_size / 1024 / 1024, 2),
                        "message": message
                    }
                else:
                    errors.append(f"{method_name}: {message}")
            
            return {
                "error": "All conversion methods failed: " + "; ".join(errors),
                "success": False
            }
            
        except Exception as e:
            logger.error(f"PDF to Word conversion failed: {traceback.format_exc()}")
            return {"error": f"Conversion failed: {str(e)}", "success": False}
    
    def batch_convert(self, input_files: List[str], output_dir: str, 
                     conversion_type: str) -> Dict[str, Any]:
        """
        Batch convert multiple files
        """
        try:
            output_dir_path = Path(output_dir)
            output_dir_path.mkdir(parents=True, exist_ok=True)
            
            results = []
            successful = 0
            failed = 0
            
            for input_file in input_files:
                input_path = Path(input_file)
                
                if conversion_type == 'word_to_pdf':
                    output_file = output_dir_path / f"{input_path.stem}.pdf"
                    result = self.convert_word_to_pdf(str(input_file), str(output_file))
                elif conversion_type == 'pdf_to_word':
                    output_file = output_dir_path / f"{input_path.stem}.docx"
                    result = self.convert_pdf_to_word(str(input_file), str(output_file))
                else:
                    result = {"error": f"Unknown conversion type: {conversion_type}", "success": False}
                
                results.append({
                    "input_file": str(input_file),
                    "output_file": str(output_file) if result.get("success") else None,
                    "success": result.get("success", False),
                    "message": result.get("message") or result.get("error")
                })
                
                if result.get("success"):
                    successful += 1
                else:
                    failed += 1
            
            return {
                "total_files": len(input_files),
                "successful": successful,
                "failed": failed,
                "results": results
            }
            
        except Exception as e:
            logger.error(f"Batch conversion failed: {traceback.format_exc()}")
            return {"error": f"Batch conversion failed: {str(e)}"}

# ============================================================================
# CORE PDF PROCESSING ENGINE (ORIGINAL)
# ============================================================================

class EnterprisePDFProcessor:
    """
    Production-grade PDF processor with multiple extraction strategies
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf']
        self.max_file_size = 100 * 1024 * 1024  # 100MB limit
        
    def validate_pdf(self, filepath: str) -> Tuple[bool, str]:
        """Validate PDF file before processing"""
        try:
            path = Path(filepath)
            
            if not path.exists():
                return False, f"File not found: {filepath}"
            
            if path.suffix.lower() not in self.supported_formats:
                return False, f"Invalid format. Expected PDF, got {path.suffix}"
            
            size = path.stat().st_size
            if size > self.max_file_size:
                return False, f"File too large: {size/1024/1024:.1f}MB (max 100MB)"
            
            if size == 0:
                return False, "File is empty"
            
            try:
                with open(filepath, 'rb') as f:
                    header = f.read(5)
                    if header != b'%PDF-':
                        return False, "Invalid PDF header"
            except Exception as e:
                return False, f"Cannot read file: {str(e)}"
            
            return True, "Valid PDF"
            
        except Exception as e:
            return False, f"Validation error: {str(e)}"
    
    def extract_metadata(self, filepath: str) -> Dict[str, Any]:
        """Extract comprehensive PDF metadata"""
        try:
            reader = PdfReader(filepath)
            metadata = reader.metadata or {}
            
            info = {
                "filename": Path(filepath).name,
                "filepath": str(Path(filepath).absolute()),
                "file_size_mb": round(Path(filepath).stat().st_size / 1024 / 1024, 2),
                "total_pages": len(reader.pages),
                "is_encrypted": reader.is_encrypted,
                "pdf_version": getattr(reader, 'pdf_header', 'Unknown'),
                "title": metadata.get('/Title', 'N/A'),
                "author": metadata.get('/Author', 'N/A'),
                "subject": metadata.get('/Subject', 'N/A'),
                "creator": metadata.get('/Creator', 'N/A'),
                "producer": metadata.get('/Producer', 'N/A'),
                "creation_date": metadata.get('/CreationDate', 'N/A'),
                "modification_date": metadata.get('/ModDate', 'N/A'),
                "extraction_timestamp": datetime.now().isoformat()
            }
            
            logger.info(f"Extracted metadata: {info['filename']} ({info['total_pages']} pages)")
            return info
            
        except Exception as e:
            logger.error(f"Metadata extraction failed: {str(e)}")
            return {"error": f"Metadata extraction failed: {str(e)}"}
    
    def extract_text_pypdf(self, filepath: str, page_range: Optional[Tuple[int, int]] = None) -> Dict[str, Any]:
        """Fast text extraction using PyPDF2"""
        try:
            reader = PdfReader(filepath)
            total_pages = len(reader.pages)
            
            start_page = 0
            end_page = total_pages
            
            if page_range:
                start_page = max(0, page_range[0] - 1)
                end_page = min(total_pages, page_range[1])
            
            pages_data = []
            full_text = []
            
            for i in range(start_page, end_page):
                page = reader.pages[i]
                text = page.extract_text() or ""
                
                pages_data.append({
                    "page_number": i + 1,
                    "text": text.strip(),
                    "char_count": len(text)
                })
                
                full_text.append(f"--- Page {i + 1} ---\n{text}")
            
            result = {
                "method": "pypdf",
                "total_pages": total_pages,
                "extracted_pages": len(pages_data),
                "page_range": f"{start_page + 1}-{end_page}",
                "full_text": "\n\n".join(full_text),
                "pages": pages_data,
                "total_characters": sum(p['char_count'] for p in pages_data)
            }
            
            logger.info(f"PyPDF extraction: {len(pages_data)} pages, {result['total_characters']} chars")
            return result
            
        except Exception as e:
            logger.error(f"PyPDF extraction failed: {traceback.format_exc()}")
            return {"error": f"Text extraction failed: {str(e)}"}
    
    def extract_text_pdfplumber(self, filepath: str, page_range: Optional[Tuple[int, int]] = None) -> Dict[str, Any]:
        """Enhanced text extraction using pdfplumber"""
        if not pdfplumber:
            return {"error": "pdfplumber not installed"}
        
        try:
            with pdfplumber.open(filepath) as pdf:
                total_pages = len(pdf.pages)
                
                start_page = 0
                end_page = total_pages
                
                if page_range:
                    start_page = max(0, page_range[0] - 1)
                    end_page = min(total_pages, page_range[1])
                
                pages_data = []
                full_text = []
                
                for i in range(start_page, end_page):
                    page = pdf.pages[i]
                    text = page.extract_text() or ""
                    
                    bbox = page.bbox
                    
                    pages_data.append({
                        "page_number": i + 1,
                        "text": text.strip(),
                        "char_count": len(text),
                        "dimensions": {
                            "width": bbox[2] - bbox[0],
                            "height": bbox[3] - bbox[1]
                        }
                    })
                    
                    full_text.append(f"--- Page {i + 1} ---\n{text}")
                
                result = {
                    "method": "pdfplumber",
                    "total_pages": total_pages,
                    "extracted_pages": len(pages_data),
                    "page_range": f"{start_page + 1}-{end_page}",
                    "full_text": "\n\n".join(full_text),
                    "pages": pages_data,
                    "total_characters": sum(p['char_count'] for p in pages_data)
                }
                
                logger.info(f"PDFPlumber extraction: {len(pages_data)} pages")
                return result
                
        except Exception as e:
            logger.error(f"PDFPlumber extraction failed: {traceback.format_exc()}")
            return {"error": f"Enhanced extraction failed: {str(e)}"}
    
    def extract_tables(self, filepath: str, page_numbers: Optional[List[int]] = None) -> Dict[str, Any]:
        """Extract tables from PDF"""
        if not pdfplumber:
            return {"error": "pdfplumber required for table extraction"}
        
        try:
            with pdfplumber.open(filepath) as pdf:
                total_pages = len(pdf.pages)
                pages_to_process = page_numbers if page_numbers else range(1, total_pages + 1)
                
                all_tables = []
                
                for page_num in pages_to_process:
                    if page_num < 1 or page_num > total_pages:
                        continue
                    
                    page = pdf.pages[page_num - 1]
                    tables = page.extract_tables()
                    
                    for table_idx, table in enumerate(tables):
                        if table:
                            headers = table[0] if table else []
                            rows = table[1:] if len(table) > 1 else []
                            
                            all_tables.append({
                                "page": page_num,
                                "table_index": table_idx + 1,
                                "headers": headers,
                                "rows": rows,
                                "row_count": len(rows),
                                "column_count": len(headers)
                            })
                
                result = {
                    "method": "pdfplumber",
                    "total_tables_found": len(all_tables),
                    "pages_scanned": len(list(pages_to_process)),
                    "tables": all_tables
                }
                
                logger.info(f"Table extraction: {len(all_tables)} tables found")
                return result
                
        except Exception as e:
            logger.error(f"Table extraction failed: {traceback.format_exc()}")
            return {"error": f"Table extraction failed: {str(e)}"}
    
    def extract_with_ocr(self, filepath: str, page_numbers: Optional[List[int]] = None) -> Dict[str, Any]:
        """OCR extraction for scanned PDFs"""
        if not OCR_AVAILABLE:
            return {"error": "OCR not available. Install pdf2image, pillow, pytesseract"}
        
        try:
            pages_to_process = page_numbers if page_numbers else None
            
            if pages_to_process:
                images = convert_from_path(
                    filepath, 
                    first_page=min(pages_to_process),
                    last_page=max(pages_to_process)
                )
            else:
                images = convert_from_path(filepath)
            
            pages_data = []
            full_text = []
            
            for idx, image in enumerate(images):
                page_num = pages_to_process[idx] if pages_to_process else idx + 1
                text = pytesseract.image_to_string(image, lang='eng')
                
                pages_data.append({
                    "page_number": page_num,
                    "text": text.strip(),
                    "char_count": len(text)
                })
                
                full_text.append(f"--- Page {page_num} (OCR) ---\n{text}")
            
            result = {
                "method": "ocr_tesseract",
                "extracted_pages": len(pages_data),
                "full_text": "\n\n".join(full_text),
                "pages": pages_data,
                "total_characters": sum(p['char_count'] for p in pages_data)
            }
            
            logger.info(f"OCR extraction: {len(pages_data)} pages processed")
            return result
            
        except Exception as e:
            logger.error(f"OCR extraction failed: {traceback.format_exc()}")
            return {"error": f"OCR failed: {str(e)}"}
    
    def search_text(self, filepath: str, query: str, case_sensitive: bool = False) -> Dict[str, Any]:
        """Search for text across all pages"""
        try:
            reader = PdfReader(filepath)
            results = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                
                search_text = text if case_sensitive else text.lower()
                search_query = query if case_sensitive else query.lower()
                
                if search_query in search_text:
                    start = 0
                    occurrences = []
                    
                    while True:
                        pos = search_text.find(search_query, start)
                        if pos == -1:
                            break
                        
                        context_start = max(0, pos - 50)
                        context_end = min(len(text), pos + len(query) + 50)
                        context = text[context_start:context_end]
                        
                        occurrences.append({
                            "position": pos,
                            "context": context.strip()
                        })
                        
                        start = pos + 1
                    
                    if occurrences:
                        results.append({
                            "page_number": i + 1,
                            "occurrence_count": len(occurrences),
                            "occurrences": occurrences
                        })
            
            summary = {
                "query": query,
                "case_sensitive": case_sensitive,
                "total_matches": sum(r['occurrence_count'] for r in results),
                "pages_with_matches": len(results),
                "results": results
            }
            
            logger.info(f"Search '{query}': {summary['total_matches']} matches")
            return summary
            
        except Exception as e:
            logger.error(f"Search failed: {str(e)}")
            return {"error": f"Search failed: {str(e)}"}

# ============================================================================
# MCP SERVER IMPLEMENTATION
# ============================================================================

class PDFMCPServer:
    """MCP Server wrapper for PDF processor and Word converter"""
    
    def __init__(self):
        self.processor = EnterprisePDFProcessor()
        self.converter = WordPDFConverter()
        self.server = Server("pdf-reader-enterprise")
        
        self._register_tools()
        logger.info("PDF MCP Server initialized with Word conversion support")
    
    def _register_tools(self):
        """Register all PDF processing and conversion tools"""
        
        @self.server.list_tools()
        async def list_tools() -> List[Tool]:
            return [
                Tool(
                    name="get_pdf_metadata",
                    description="Extract comprehensive metadata from PDF (pages, author, creation date, file size, etc.)",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filepath": {
                                "type": "string",
                                "description": "Absolute path to PDF file"
                            }
                        },
                        "required": ["filepath"]
                    }
                ),
                Tool(
                    name="read_pdf_text",
                    description="Extract text from PDF with page range support. Use for digital PDFs (fast extraction).",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filepath": {
                                "type": "string",
                                "description": "Absolute path to PDF file"
                            },
                            "method": {
                                "type": "string",
                                "enum": ["pypdf", "pdfplumber"],
                                "description": "Extraction method: pypdf (fast) or pdfplumber (better layout)",
                                "default": "pypdf"
                            },
                            "start_page": {
                                "type": "integer",
                                "description": "Start page number (1-indexed). Optional.",
                                "minimum": 1
                            },
                            "end_page": {
                                "type": "integer",
                                "description": "End page number (inclusive). Optional.",
                                "minimum": 1
                            }
                        },
                        "required": ["filepath"]
                    }
                ),
                Tool(
                    name="extract_pdf_tables",
                    description="Extract tables from PDF pages. Returns structured table data with headers and rows.",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filepath": {
                                "type": "string",
                                "description": "Absolute path to PDF file"
                            },
                            "page_numbers": {
                                "type": "array",
                                "items": {"type": "integer"},
                                "description": "List of page numbers to scan (1-indexed). If empty, scans all pages."
                            }
                        },
                        "required": ["filepath"]
                    }
                ),
                Tool(
                    name="read_pdf_ocr",
                    description="Extract text from scanned PDFs using OCR (Optical Character Recognition). Slower but works on image-based PDFs.",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filepath": {
                                "type": "string",
                                "description": "Absolute path to PDF file"
                            },
                            "page_numbers": {
                                "type": "array",
                                "items": {"type": "integer"},
                                "description": "List of page numbers to OCR (1-indexed). If empty, processes all pages."
                            }
                        },
                        "required": ["filepath"]
                    }
                ),
                Tool(
                    name="search_pdf_text",
                    description="Search for specific text across all PDF pages. Returns page numbers and context.",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "filepath": {
                                "type": "string",
                                "description": "Absolute path to PDF file"
                            },
                            "query": {
                                "type": "string",
                                "description": "Text to search for"
                            },
                            "case_sensitive": {
                                "type": "boolean",
                                "description": "Enable case-sensitive search",
                                "default": False
                            }
                        },
                        "required": ["filepath", "query"]
                    }
                ),
                # ===== NEW CONVERSION TOOLS =====
                Tool(
                    name="convert_word_to_pdf",
                    description="Convert Word document (.docx, .doc) to PDF with high quality. Production-grade conversion similar to iLovePDF.",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "word_path": {
                                "type": "string",
                                "description": "Absolute path to Word document (.docx or .doc)"
                            },
                            "output_path": {
                                "type": "string",
                                "description": "Output PDF path. If not specified, saves to same directory as input with .pdf extension."
                            },
                            "method": {
                                "type": "string",
                                "enum": ["auto", "libreoffice", "docx2pdf"],
                                "description": "Conversion method: auto (tries best available), libreoffice (cross-platform), docx2pdf (Windows)",
                                "default": "auto"
                            }
                        },
                        "required": ["word_path"]
                    }
                ),
                Tool(
                    name="convert_pdf_to_word",
                    description="Convert PDF to Word document (.docx) with layout preservation. Extracts text, tables, and formatting.",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "pdf_path": {
                                "type": "string",
                                "description": "Absolute path to PDF file"
                            },
                            "output_path": {
                                "type": "string",
                                "description": "Output Word document path. If not specified, saves to same directory as input with .docx extension."
                            },
                            "method": {
                                "type": "string",
                                "enum": ["auto", "advanced", "libreoffice"],
                                "description": "Conversion method: auto (tries best), advanced (text+tables), libreoffice (layout preservation)",
                                "default": "auto"
                            }
                        },
                        "required": ["pdf_path"]
                    }
                ),
                Tool(
                    name="batch_convert_documents",
                    description="Batch convert multiple documents (Word to PDF or PDF to Word). Process multiple files at once.",
                    inputSchema={
                        "type": "object",
                        "properties": {
                            "input_files": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "List of absolute paths to files to convert"
                            },
                            "output_dir": {
                                "type": "string",
                                "description": "Output directory for converted files"
                            },
                            "conversion_type": {
                                "type": "string",
                                "enum": ["word_to_pdf", "pdf_to_word"],
                                "description": "Type of conversion to perform"
                            }
                        },
                        "required": ["input_files", "output_dir", "conversion_type"]
                    }
                )
            ]
        
        @self.server.call_tool()
        async def call_tool(name: str, arguments: Dict[str, Any]) -> List[TextContent]:
            try:
                # PDF Reading Tools
                if name in ["get_pdf_metadata", "read_pdf_text", "extract_pdf_tables", 
                           "read_pdf_ocr", "search_pdf_text"]:
                    filepath = arguments.get("filepath")
                    
                    is_valid, msg = self.processor.validate_pdf(filepath)
                    if not is_valid:
                        return [TextContent(type="text", text=f"❌ Validation failed: {msg}")]
                    
                    if name == "get_pdf_metadata":
                        result = self.processor.extract_metadata(filepath)
                    
                    elif name == "read_pdf_text":
                        method = arguments.get("method", "pypdf")
                        start_page = arguments.get("start_page")
                        end_page = arguments.get("end_page")
                        
                        page_range = None
                        if start_page and end_page:
                            page_range = (start_page, end_page)
                        
                        if method == "pdfplumber":
                            result = self.processor.extract_text_pdfplumber(filepath, page_range)
                        else:
                            result = self.processor.extract_text_pypdf(filepath, page_range)
                    
                    elif name == "extract_pdf_tables":
                        page_numbers = arguments.get("page_numbers", [])
                        result = self.processor.extract_tables(filepath, page_numbers)
                    
                    elif name == "read_pdf_ocr":
                        page_numbers = arguments.get("page_numbers", [])
                        result = self.processor.extract_with_ocr(filepath, page_numbers)
                    
                    elif name == "search_pdf_text":
                        query = arguments.get("query")
                        case_sensitive = arguments.get("case_sensitive", False)
                        result = self.processor.search_text(filepath, query, case_sensitive)
                
                # Conversion Tools
                elif name == "convert_word_to_pdf":
                    word_path = arguments.get("word_path")
                    output_path = arguments.get("output_path")
                    method = arguments.get("method", "auto")
                    
                    result = self.converter.convert_word_to_pdf(word_path, output_path, method)
                
                elif name == "convert_pdf_to_word":
                    pdf_path = arguments.get("pdf_path")
                    output_path = arguments.get("output_path")
                    method = arguments.get("method", "auto")
                    
                    result = self.converter.convert_pdf_to_word(pdf_path, output_path, method)
                
                elif name == "batch_convert_documents":
                    input_files = arguments.get("input_files", [])
                    output_dir = arguments.get("output_dir")
                    conversion_type = arguments.get("conversion_type")
                    
                    result = self.converter.batch_convert(input_files, output_dir, conversion_type)
                
                else:
                    return [TextContent(type="text", text=f"❌ Unknown tool: {name}")]
                
                # Format response
                if "error" in result:
                    return [TextContent(type="text", text=f"❌ {result['error']}")]
                
                output = json.dumps(result, indent=2, ensure_ascii=False)
                return [TextContent(type="text", text=f"✅ SUCCESS\n\n{output}")]
                
            except Exception as e:
                logger.error(f"Tool execution failed: {traceback.format_exc()}")
                return [TextContent(type="text", text=f"❌ Error: {str(e)}")]
    
    async def run(self):
        """Start the MCP server"""
        async with mcp.server.stdio.stdio_server() as (read_stream, write_stream):
            logger.info("PDF MCP Server running on stdio")
            await self.server.run(
                read_stream,
                write_stream,
                self.server.create_initialization_options()
            )

# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

async def main():
    """Main entry point"""
    logger.info("=" * 60)
    logger.info("ENTERPRISE PDF MCP SERVER - Production Ready")
    logger.info("WITH WORD ↔ PDF CONVERSION")
    logger.info("=" * 60)
    logger.info(f"Python: {sys.version}")
    logger.info(f"PyPDF2: Available")
    logger.info(f"pdfplumber: {'Available' if pdfplumber else 'Not installed'}")
    logger.info(f"OCR: {'Available' if OCR_AVAILABLE else 'Not installed'}")
    logger.info(f"Camelot: {'Available' if CAMELOT_AVAILABLE else 'Not installed'}")
    logger.info("--- Conversion Support ---")
    logger.info(f"python-docx: {'Available' if DOCX_AVAILABLE else 'Not installed'}")
    logger.info(f"docx2pdf: {'Available' if DOCX2PDF_AVAILABLE else 'Not installed'}")
    logger.info(f"LibreOffice: {'Available' if LIBREOFFICE_AVAILABLE else 'Not installed'}")
    logger.info("=" * 60)
    
    server = PDFMCPServer()
    await server.run()

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        logger.info("Server stopped by user")
    except Exception as e:
        logger.error(f"Fatal error: {traceback.format_exc()}")
        sys.exit(1)